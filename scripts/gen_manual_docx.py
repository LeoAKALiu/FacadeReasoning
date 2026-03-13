#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate 软件使用说明书.docx from rz_polished.md.
Adds header, footer, heading levels, body text, and figure placeholders.
Run from project root: python3 scripts/gen_manual_docx.py
"""
from __future__ import print_function

import re
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Please install: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# Paths (run from project root)
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = os.path.join(PROJECT_ROOT, "rz_polished.md")
OUT_PATH = os.path.join(PROJECT_ROOT, "软件使用说明书.docx")
SCREENSHOTS_DIR = os.path.join(PROJECT_ROOT, "screenshots")

HEADER_TEXT = "建筑立面参数化推理与可靠度评估系统 V1.0"


def set_cell_shading(cell, color):
    """Set table cell shading (e.g. for placeholder)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def add_page_number_footer(section):
    """Add centered page number to section footer."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)


def set_header_footer(doc):
    """Set same header text and page number footer for all sections."""
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        r = p.add_run(HEADER_TEXT)
        r.font.size = Pt(9)
        r.font.name = "宋体"
        add_page_number_footer(section)


def parse_md(path):
    """Read markdown and yield (block_type, content). block_type: h1, h2, h3, p, li, img, hr, skip."""
    with open(path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # Skip (页眉)(页码) metadata lines
        if re.match(r"^\*\*\(页眉", stripped) or re.match(r"^\*\*\(页码", stripped):
            i += 1
            continue
        if stripped.startswith("## ") and not stripped.startswith("### "):
            yield ("h1", stripped[3:].strip())
            i += 1
            continue
        if stripped.startswith("### "):
            yield ("h2", stripped[4:].strip())
            i += 1
            continue
        if stripped.startswith("#### "):
            yield ("h3", stripped[5:].strip())
            i += 1
            continue
        if stripped == "---":
            yield ("hr", "")
            i += 1
            continue
        # Image: ![图N：...](图N.png)
        img_match = re.match(r"^!\[(图\d+：[^\]]+)\]\(([^)]+)\)", stripped)
        if img_match:
            yield ("img", (img_match.group(1), img_match.group(2)))
            i += 1
            continue
        # List item: "- " or "1. "
        if stripped.startswith("- ") or re.match(r"^\d+\.\s", stripped):
            yield ("li", stripped)
            i += 1
            continue
        # Bold list like "**4.1.1 按钮与链接**"
        if stripped.startswith("**") and stripped.endswith("**") and " " in stripped:
            yield ("p", stripped)
            i += 1
            continue
        # Empty line
        if not stripped:
            i += 1
            continue
        # Paragraph (possibly multi-line until blank or next heading/list)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            next_stripped = next_line.strip()
            if not next_stripped:
                break
            if next_stripped.startswith("## ") or next_stripped.startswith("### ") or next_stripped.startswith("- ") or next_stripped.startswith("!["):
                break
            if re.match(r"^\d+\.\s", next_stripped):
                break
            para_lines.append(next_stripped)
            i += 1
        yield ("p", "\n".join(para_lines))
    return


def add_paragraph(doc, text, style=None):
    """Add paragraph; strip ** for bold markers and apply style."""
    p = doc.add_paragraph()
    if style:
        p.style = style
    # Simple bold: **text** -> bold run
    rest = text
    while rest:
        m = re.search(r"\*\*([^*]+)\*\*", rest)
        if not m:
            p.add_run(rest)
            break
        p.add_run(rest[: m.start()])
        r = p.add_run(m.group(1))
        r.bold = True
        rest = rest[m.end() :]
    return p


def main():
    os.chdir(PROJECT_ROOT)
    if not os.path.isfile(MD_PATH):
        print("Not found: %s" % MD_PATH, file=sys.stderr)
        sys.exit(1)

    doc = Document()
    # Set default font for body
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for block_type, content in parse_md(MD_PATH):
        if block_type == "h1":
            p = doc.add_paragraph(content)
            p.style = "Heading 1"
            p.runs[0].font.name = "黑体"
            p.runs[0].font.size = Pt(16)
        elif block_type == "h2":
            p = doc.add_paragraph(content)
            p.style = "Heading 2"
            p.runs[0].font.name = "黑体"
            p.runs[0].font.size = Pt(14)
        elif block_type == "h3":
            p = doc.add_paragraph(content)
            p.style = "Heading 3"
            p.runs[0].font.name = "黑体"
            p.runs[0].font.size = Pt(12)
        elif block_type == "hr":
            doc.add_paragraph()
        elif block_type == "li":
            p = doc.add_paragraph()
            p.style = "List Bullet"
            t = content.strip()
            if t.startswith("- "):
                t = t[2:]
            elif re.match(r"^\d+\.\s", t):
                t = re.sub(r"^\d+\.\s", "", t)
            for part in re.split(r"(\*\*[^*]+\*\*)", t):
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)
        elif block_type == "img":
            caption, filename = content
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption)
            r.font.size = Pt(10.5)
            r.font.name = "宋体"
            image_path = os.path.join(SCREENSHOTS_DIR, filename)
            if os.path.isfile(image_path):
                doc.add_paragraph()
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_img.add_run()
                run.add_picture(image_path, width=Cm(14))
                doc.add_paragraph()
            else:
                # Placeholder when screenshot not found
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False
                cell = table.rows[0].cells[0]
                cell.width = Cm(12)
                cell.text = "【请在此插入%s】" % caption
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = None
                set_cell_shading(cell, "E7E6E6")
                doc.add_paragraph()
        elif block_type == "p":
            if not content.strip():
                continue
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*[^*]+\*\*)", content):
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)

    set_header_footer(doc)
    doc.save(OUT_PATH)
    print("Generated: %s" % OUT_PATH)


if __name__ == "__main__":
    main()